from docx import Document

def generate_docx(cv_data, output_path):
    doc = Document()
    doc.add_heading(cv_data.get("name", ""), 0)
    doc.add_paragraph(f"Email: {cv_data.get('email', '')}")
    doc.add_heading("Contenu du CV", level=1)
    doc.add_paragraph(cv_data.get("content", ""))
    doc.save(output_path)