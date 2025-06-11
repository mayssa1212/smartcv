from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import Dict, Optional
from pydantic import BaseModel
import os
import tempfile
from docx import Document
import pdfkit

from app.database import get_db
from app.auth import get_current_user
from app import models
from app import nlp_utils

router = APIRouter(
    prefix="/generator",
    tags=["generator"]
)

class CVGenerationRequest(BaseModel):
    cv_id: int
    template: str = "modern"
    color_theme: str = "blue"
    font: str = "arial"

class CoverLetterRequest(BaseModel):
    job_title: str
    company_name: str
    cv_id: int
    tone: str = "professional"  # professional, friendly, enthusiastic

@router.post("/generate-pdf")
def generate_pdf(
    request: CVGenerationRequest, 
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_current_user)
):
    """
    Génère un CV au format PDF
    """
    # Récupérer le CV depuis la base de données
    cv = db.query(models.CV).filter(models.CV.id == request.cv_id, models.CV.user_id == current_user.id).first()
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")
    
    # Créer un fichier HTML temporaire
    with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as temp_html:
        html_content = generate_html_from_cv(cv, request.template, request.color_theme, request.font)
        temp_html.write(html_content.encode('utf-8'))
        html_path = temp_html.name
    
    # Créer un fichier PDF temporaire
    pdf_path = f"{html_path}.pdf"
    
    # Générer le PDF à partir du HTML
    try:
        pdfkit.from_file(html_path, pdf_path)
    except Exception as e:
        # Nettoyer les fichiers temporaires
        os.unlink(html_path)
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)
        raise HTTPException(status_code=500, detail=f"PDF generation error: {str(e)}")
    
    # Nettoyer le fichier HTML temporaire
    background_tasks.add_task(os.unlink, html_path)
    
    # Nettoyer le fichier PDF après l'envoi
    background_tasks.add_task(os.unlink, pdf_path)
    
    # Renvoyer le fichier PDF
    return FileResponse(
        path=pdf_path,
        filename=f"cv_{current_user.id}_{cv.id}.pdf",
        media_type="application/pdf",
        background=background_tasks
    )

@router.post("/generate-docx")
def generate_docx(
    request: CVGenerationRequest, 
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_current_user)
):
    """
    Génère un CV au format DOCX
    """
    # Récupérer le CV depuis la base de données
    cv = db.query(models.CV).filter(models.CV.id == request.cv_id, models.CV.user_id == current_user.id).first()
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")
    
    # Créer un document Word
    doc = Document()
    
    # Ajouter les informations du CV au document
    doc.add_heading(f"{cv.personal_info.get('full_name', 'CV')}", 0)
    
    # Informations personnelles
    doc.add_heading("Informations personnelles", level=1)
    for key, value in cv.personal_info.items():
        if key != "full_name":  # Déjà utilisé comme titre
            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
    
    # Formation
    if cv.education:
        doc.add_heading("Formation", level=1)
        for edu in cv.education:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')}, {edu.get('institution', '')}").bold = True
            p.add_run(f"\n{edu.get('start_date', '')} - {edu.get('end_date', '')}")
            p.add_run(f"\n{edu.get('description', '')}")
    
    # Expérience professionnelle
    if cv.experience:
        doc.add_heading("Expérience professionnelle", level=1)
        for exp in cv.experience:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('position', '')}, {exp.get('company', '')}").bold = True
            p.add_run(f"\n{exp.get('start_date', '')} - {exp.get('end_date', '')}")
            p.add_run(f"\n{exp.get('description', '')}")
    
    # Compétences
    if cv.skills:
        doc.add_heading("Compétences", level=1)
        for skill in cv.skills:
            doc.add_paragraph(skill, style="List Bullet")
    
    # Créer un fichier temporaire pour le document Word
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
        docx_path = temp_docx.name
        doc.save(docx_path)
    
    # Nettoyer le fichier après l'envoi
    background_tasks.add_task(os.unlink, docx_path)
    
    # Renvoyer le fichier DOCX
    return FileResponse(
        path=docx_path,
        filename=f"cv_{current_user.id}_{cv.id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=background_tasks
    )

def generate_html_from_cv(cv, template, color_theme, font):
    """
    Génère le contenu HTML pour un CV selon le template choisi
    """
    # Styles CSS de base
    css = f"""
    body {{
        font-family: {font}, sans-serif;
        margin: 0;
        padding: 0;
        color: #333;
    }}
    """
    
    # Ajouter des styles spécifiques selon le thème de couleur
    if color_theme == "blue":
        css += """
        h1, h2, h3 { color: #2c3e50; }
        .header { background-color: #3498db; color: white; }
        """
    elif color_theme == "green":
        css += """
        h1, h2, h3 { color: #27ae60; }
        .header { background-color: #2ecc71; color: white; }
        """
    elif color_theme == "purple":
        css += """
        h1, h2, h3 { color: #8e44ad; }
        .header { background-color: #9b59b6; color: white; }
        """
    
    # Contenu HTML selon le template
    if template == "modern":
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>CV - {cv.personal_info.get('full_name', '')}</title>
            <style>{css}</style>
        </head>
        <body>
            <div class="header" style="padding: 20px;">
                <h1>{cv.personal_info.get('full_name', '')}</h1>
                <p>{cv.personal_info.get('email', '')}</p>
                <p>{cv.personal_info.get('phone', '')}</p>
            </div>
            
            <div style="padding: 20px;">
                <h2>Formation</h2>
                {''.join([f'<div style="margin-bottom: 15px;"><strong>{edu.get("degree", "")}, {edu.get("institution", "")}</strong><br>{edu.get("start_date", "")} - {edu.get("end_date", "")}<br>{edu.get("description", "")}</div>' for edu in cv.education])}
                
                <h2>Expérience professionnelle</h2>
                {''.join([f'<div style="margin-bottom: 15px;"><strong>{exp.get("position", "")}, {exp.get("company", "")}</strong><br>{exp.get("start_date", "")} - {exp.get("end_date", "")}<br>{exp.get("description", "")}</div>' for exp in cv.experience])}
                
                <h2>Compétences</h2>
                <ul>
                    {''.join([f'<li>{skill}</li>' for skill in cv.skills])}
                </ul>
            </div>
        </body>
        </html>
        """
    elif template == "classic":
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>CV - {cv.personal_info.get('full_name', '')}</title>
            <style>{css}</style>
        </head>
        <body>
            <div style="padding: 20px;">
                <h1 style="text-align: center;">{cv.personal_info.get('full_name', '')}</h1>
                <p style="text-align: center;">{cv.personal_info.get('email', '')} | {cv.personal_info.get('phone', '')}</p>
                
                <hr>
                
                <h2>Formation</h2>
                {''.join([f'<div style="margin-bottom: 15px;"><strong>{edu.get("degree", "")}</strong>, {edu.get("institution", "")}<br>{edu.get("start_date", "")} - {edu.get("end_date", "")}<br>{edu.get("description", "")}</div>' for edu in cv.education])}
                
                <h2>Expérience professionnelle</h2>
                {''.join([f'<div style="margin-bottom: 15px;"><strong>{exp.get("position", "")}</strong>, {exp.get("company", "")}<br>{exp.get("start_date", "")} - {exp.get("end_date", "")}<br>{exp.get("description", "")}</div>' for exp in cv.experience])}
                
                <h2>Compétences</h2>
                <ul>
                    {''.join([f'<li>{skill}</li>' for skill in cv.skills])}
                </ul>
            </div>
        </body>
        </html>
        """
    else:
        # Template par défaut si aucun des templates connus n'est spécifié
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>CV - {cv.personal_info.get('full_name', '')}</title>
            <style>{css}</style>
        </head>
        <body>
            <h1>{cv.personal_info.get('full_name', '')}</h1>
            <p>{cv.personal_info.get('email', '')}</p>
            <p>{cv.personal_info.get('phone', '')}</p>
            
            <h2>Formation</h2>
            {''.join([f'<div><strong>{edu.get("degree", "")}</strong>, {edu.get("institution", "")}<br>{edu.get("start_date", "")} - {edu.get("end_date", "")}</div>' for edu in cv.education])}
            
            <h2>Expérience</h2>
            {''.join([f'<div><strong>{exp.get("position", "")}</strong>, {exp.get("company", "")}<br>{exp.get("start_date", "")} - {exp.get("end_date", "")}</div>' for exp in cv.experience])}
            
            <h2>Compétences</h2>
            <ul>
                {''.join([f'<li>{skill}</li>' for skill in cv.skills])}
            </ul>
        </body>
        </html>
        """
    
    return html

@router.post("/generate-cover-letter")
def generate_cover_letter(
    request: CoverLetterRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Génère une lettre de motivation adaptée au poste visé
    """
    # Récupérer le CV
    cv = db.query(models.CV).filter(models.CV.id == request.cv_id, models.CV.user_id == current_user.id).first()
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")
    
    # Générer la lettre de motivation avec l'IA
    cover_letter = nlp_utils.generate_cover_letter(
        cv_text=cv.data,
        job_title=request.job_title,
        company_name=request.company_name,
        tone=request.tone
    )
    
    return {"cover_letter": cover_letter}

